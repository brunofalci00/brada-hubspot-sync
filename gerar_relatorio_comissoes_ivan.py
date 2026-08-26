# -*- coding: utf-8 -*-
"""Gera o Relatorio de Vendas e Comissoes (MD + DOCX + PDF) na identidade Brada.

Le as abas Junho_MATCH e Junho_Elaboracao de Projetos da planilha Comissoes 2026
(que ja puxa do HubSpot) e monta um relatorio para o financeiro com:
  - Sumario executivo (narrativa, sem grafico).
  - Visao consolidada (subtotais: interno/externo, esfera de incentivo, lei).
  - Detalhamento MATCH e Elaboracao (colunas de comissao/pagamento em branco).
As comissoes ficam em branco de proposito (o financeiro preenche na planilha de folha).

Estilo Brada: Calibri, laranja C55A11, sem tracos como separador, sem grafico, sem logo.
Rota PDF: HTML -> navegador headless (--print-to-pdf), dialog-free e portavel.

Vive no repo desde 20/08. Antes ficava solto em HubSpot/Scripts/, fora de git: um documento que
vai ao financeiro sem historico e sem backup. Ao entrar, trocou as tres amarras de maquina por
mecanismos que o resto do repo ja usa: credencial pelo get_sheets_client (cascata que serve dev e
CI), navegador por descoberta em vez de caminho fixo do Windows, e pasta de saida por --out/env.

Uso:
  python gerar_relatorio_comissoes_ivan.py --cycle 2026-08 --emissao 2026-08-20
"""
import argparse
import datetime
import html as _htmllib
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from gspread.utils import rowcol_to_a1
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from sync import get_sheets_client                                     # noqa: E402
from sheets_comissoes_ivan import OFICIAL_ID_DEFAULT, utf8_stdout      # noqa: E402

# ------------------------------------------------------------------ config
SID = OFICIAL_ID_DEFAULT
MATCH_TAB = "Junho_MATCH"
ELAB_TAB = "Junho_Elaboração de Projetos"  # aba original (gid 1005722547), nao a "Cópia de"
# Pasta de saida. O default e a que o Bruno ja usa, pra nao mudar o fluxo dele; em outra
# maquina ou no runner, passa --out ou define RELATORIOS_OUT.
DEFAULT_OUT = (os.environ.get("RELATORIOS_OUT")
               or r"C:\Users\bruno\Documents\Brada\HubSpot\Relatorios_Financeiro")
# Defaults = ciclo de Junho. Sobrescritos por --cycle via cycle_labels().
#
# O rotulo diz 21 e nao 20 porque a janela real do fecho e 21 do mes anterior ate 20 deste
# (cycle_window, com o teste de fronteira em ops/_test_comissoes_ivan.py: 20/07 fora, 21/07
# dentro). Escrever "20/06 a 20/07" e depois "20/07 a 20/08" faz o dia 20/07 aparecer no rotulo
# de dois relatorios seguidos, o que num documento de apuracao le como periodo sobreposto.
# O NOME DO ARQUIVO fica com 20 nos dois lados, pra nao quebrar a serie ja entregue.
BASENAME = "Relatorio_Vendas_Comissoes_Brada_ciclo_20mai_20jun_2026"
CICLO = "Ciclo de 21/05/2026 a 20/06/2026"
PERIODO = "21/05 a 20/06"  # fragmento curto usado no sumario executivo

MES_PT = {1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho",
          7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"}
WIN_CHROME_1 = r"C:\Program Files\Google\Chrome\Application\chrome.exe"
WIN_CHROME_2 = r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe"

MES_ABBR = {1: "jan", 2: "fev", 3: "mar", 4: "abr", 5: "mai", 6: "jun",
            7: "jul", 8: "ago", 9: "set", 10: "out", 11: "nov", 12: "dez"}


def cycle_labels(cycle):
    """'2026-07' -> dict(match_tab, elab_tab, ciclo, periodo, basename). O ciclo
    YYYY-MM fecha no dia 20 do mes MM e cobre 20/(MM-1) a 20/MM (nome do mes = MM)."""
    y, m = int(cycle[:4]), int(cycle[5:7])
    pm = 12 if m == 1 else m - 1
    py = y - 1 if m == 1 else y
    mes = MES_PT[m]
    return {
        "match_tab": f"{mes}_MATCH",
        "elab_tab": f"{mes}_Elaboração de Projetos",
        "ciclo": f"Ciclo de 21/{pm:02d}/{py} a 20/{m:02d}/{y}",
        "periodo": f"21/{pm:02d} a 20/{m:02d}",
        "basename": f"Relatorio_Vendas_Comissoes_Brada_ciclo_20{MES_ABBR[pm]}_20{MES_ABBR[m]}_{y}",
    }

# Nota que substitui as tabelas de MATCH quando o ciclo fecha zerado. A frase existe porque em
# 16/07 a ausencia foi lida como falha do relatorio: havia matches em andamento no funil, e
# "match em andamento" nao e "match ganho". Sem a explicacao a mesma pergunta volta todo ciclo
# que fecha em zero, e agosto e o segundo seguido. Uma constante, nao tres copias: a frase
# aparece nos tres renderizadores e tem que sair igual nos tres.
NOTA_MATCH_VAZIO = ("Não houve conversão de MATCH no período. Este relatório considera apenas o "
                    "match com aporte confirmado; negócio de match em andamento no funil não "
                    "entra na apuração.")

# paleta Brada
ORANGE = RGBColor(0xC5, 0x5A, 0x11)
DARK = RGBColor(0x33, 0x33, 0x33)
LIGHT = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "C55A11"
ROW_ALT_BG = "FBE5D6"
TOTAL_BG = "F2E2D6"

L, C, R = None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT

# MATCH: 16 colunas. 8 de dados + 3 de contato do proponente (Nome/Telefone/E-mail,
# lidas da planilha; pedido do Ivan 10/07) + 5 comissoes (Ivan..Rafaela) em branco.
MATCH_HEADERS = ["Cliente", "Fonte de recurso", "Proponente", "Interno/Externo",
                 "Projeto", "Numero do projeto", "Valor", "Data do aporte",
                 "Nome do contato", "Telefone do proponente", "E-mail do proponente",
                 "Ivan", "Jaqueline", "Carina", "Danielle", "Rafaela"]
MATCH_TOTAL_IDX = 6            # coluna Valor (soma pro resumo)
MATCH_VALOR_IDX = 6           # idem, usado nas agregacoes
MATCH_MONEY = {6}             # colunas monetarias (prefixo R$)
MATCH_BLANK = {11, 12, 13, 14, 15}  # comissoes: financeiro preenche
MATCH_FONTE_IDX = 1
MATCH_IE_IDX = 3
MATCH_WIDTHS = [2.0, 2.0, 2.3, 1.3, 2.2, 2.1, 1.6, 1.2, 1.6, 1.9, 2.8, 0.9, 1.0, 0.9, 1.0, 0.9]
MATCH_ALIGN = [L, L, L, C, L, L, R, C, L, L, L, R, R, R, R, R]

# Tabela "MATCH externo (para cobranca)": so os externos, colunas de cobranca (sem comissao).
# Reaproveita as linhas de exibicao do MATCH, selecionando este subconjunto de colunas.
COBR_SRC_IDX = [0, 2, 4, 5, 6, 7, 8, 9, 10]
COBR_HEADERS = ["Cliente", "Proponente", "Projeto", "Numero do projeto", "Valor",
                "Data do aporte", "Nome do contato", "Telefone", "E-mail"]
COBR_WIDTHS = [2.6, 2.9, 3.0, 2.8, 1.8, 1.6, 1.9, 2.2, 3.2]
COBR_ALIGN = [L, L, L, L, R, C, L, L, L]


def cobranca_rows(match_rows):
    """Linhas externas do MATCH, projetadas nas colunas de cobranca."""
    out = []
    for r in match_rows:
        ie = (r[MATCH_IE_IDX] if len(r) > MATCH_IE_IDX else "").strip().lower()
        if ie == "externo":
            out.append([r[i] if i < len(r) else "" for i in COBR_SRC_IDX])
    return out

# Elaboracao: 12 colunas. Mostra dados ate Valor Pago (Data de pagamento + Valor Pago vem da
# origem, ja preenchidos pelo financeiro nos quitados; pedido do Ivan 07/07). OBS/Liquido/
# Ivan/Ricardo ficam em branco (financeiro preenche).
ELAB_HEADERS = ["Nome do Proponente", "Data do fechamento", "Tipo venda",
                "Condição de Pagamento", "Valor", "Lei da Submissão",
                "Data de pagamento", "Valor Pago", "OBS", "Líquido pago", "Ivan", "Ricardo"]
ELAB_TOTAL_IDX = 4            # coluna Valor (soma "fechado" pro resumo)
ELAB_VALOR_IDX = 4
ELAB_MONEY = {4, 7}          # Valor + Valor Pago recebem R$
ELAB_BLANK = {8, 9, 10, 11}  # OBS, Liquido pago, Ivan, Ricardo em branco
ELAB_LEI_IDX = 5
ELAB_COND_IDX = 3            # coluna "Condição de Pagamento" (define "a apurar" = captação)
ELAB_WIDTHS = [4.0, 1.8, 1.8, 2.6, 1.8, 2.8, 1.8, 1.6, 1.6, 1.6, 1.0, 1.2]
ELAB_ALIGN = [L, C, L, L, R, L, C, R, L, R, R, R]


# ------------------------------------------------------------------ helpers de dado
def parse_brl(s):
    s = (s or "").strip().replace("R$", "").strip()
    if not s:
        return 0.0
    try:
        return float(s.replace(".", "").replace(",", "."))
    except ValueError:
        return 0.0


def fmt_brl(v):
    return f"{v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def match_esfera(fonte):
    """Classifica a fonte de recurso do MATCH por esfera de incentivo."""
    f = (fonte or "").upper()
    if "ICMS" in f:
        return "ICMS (estadual)"
    if f.startswith("IR ") or f == "IR" or "IMPOSTO DE RENDA" in f:
        return "IR (federal)"
    if "ROUANET" in f or "AUDIOVISUAL" in f or "PRONAC" in f:
        return "Rouanet/Audiovisual (federal)"
    if "PROMAC" in f or "ISS" in f:
        return "ISS/PROMAC (municipal)"
    return "Outras fontes"


def norm_header(s):
    return " ".join((s or "").split()).casefold()


def validar_header(lidos, headers, blank_cols, tab):
    """Aborta se o cabecalho da aba nao for o que o relatorio espera.

    Sem isto, `read_tab` fatia por POSICAO sem conferir nada: uma coluna inserida no meio faz o
    PDF reportar a coluna errada ao financeiro, calado. Em 19/08 estas mesmas abas ganharam
    coluna nova; foi no fim e nao quebrou, mas foi sorte de posicao.

    Duas ressalvas, as duas vindas da planilha real e nao da teoria:

    1. Compara NORMALIZADO. Varios cabecalhos do MATCH carregam espaco a direita na planilha
       ("Proponente ", "Valor ", "Data do aporte ", "Numero do projeto "). Comparacao literal
       abortaria hoje mesmo.
    2. Confere so as colunas que o relatorio LE. As de comissao saem em branco de proposito, e
       sao justamente as que mudam de rotulo: em Julho_Elaboracao o Ivan ja renomeou "Ivan" pra
       "Ivan 5%". Travar onde nao ha leitura so produziria alarme falso.
    """
    faltas = []
    for i, esperado in enumerate(headers):
        if i in blank_cols:
            continue
        achado = lidos[i] if i < len(lidos) else ""
        if norm_header(achado) != norm_header(esperado):
            col = rowcol_to_a1(1, i + 1).rstrip("1")
            faltas.append(f"{col} (col {i}): esperado {esperado!r}, achado {achado!r}")
    if faltas:
        cabeca = f"[abort] cabecalho de {tab!r} divergente do que o relatorio le:"
        rodape = "Conferir a planilha antes de gerar: o PDF vai pro financeiro."
        raise SystemExit("\n  ".join([cabeca] + faltas + [rodape]))


def read_tab(sh, tab, headers, blank_cols, money_cols, total_idx):
    """Le a aba e devolve (linhas de exibicao, linhas cruas, soma total, nº total vazio).

    Exibicao: colunas copiadas da origem, exceto blank_cols (forcadas vazias);
    money_cols recebem prefixo R$. Cruas: celulas originais (strip) pra agregacoes.
    total_idx: coluna somada pro resumo (Valor).
    """
    ws = sh.worksheet(tab)
    vals = ws.get_all_values()
    validar_header(vals[0] if vals else [], headers, blank_cols, tab)
    disp_out, raw_out = [], []
    valor_total, n_valor_vazio = 0.0, 0
    for row in vals[1:]:
        first = (row[0] if row else "").strip()
        if not first or first.lower() == "total":
            continue
        raw = [(row[i].strip() if i < len(row) else "") for i in range(len(headers))]
        disp = []
        for i in range(len(headers)):
            if i in blank_cols:
                disp.append("")
                continue
            cell = raw[i]
            if i in money_cols and cell:
                cell = "R$ " + cell
            disp.append(cell)
        valor_total += parse_brl(raw[total_idx])
        if not raw[total_idx]:
            n_valor_vazio += 1
        disp_out.append(disp)
        raw_out.append(raw)
    return disp_out, raw_out, valor_total, n_valor_vazio


def build_summary(match_raw, elab_raw, match_total, elab_total, elab_apurar, periodo=PERIODO):
    """Agrega subtotais e monta o texto do sumario executivo."""
    ie = {}
    for r in match_raw:
        k = (r[MATCH_IE_IDX] or "").strip() or "Não classificado"
        c, s = ie.get(k, (0, 0.0))
        ie[k] = (c + 1, s + parse_brl(r[MATCH_VALOR_IDX]))
    esf = {}
    for r in match_raw:
        k = match_esfera(r[MATCH_FONTE_IDX])
        c, s = esf.get(k, (0, 0.0))
        esf[k] = (c + 1, s + parse_brl(r[MATCH_VALOR_IDX]))
    lei = {}
    for r in elab_raw:
        k = (r[ELAB_LEI_IDX] or "").strip() or "Sem lei informada"
        lei[k] = lei.get(k, 0) + 1

    ie_rows = sorted(ie.items(), key=lambda kv: -kv[1][1])
    esf_rows = sorted(esf.items(), key=lambda kv: -kv[1][1])
    lei_rows = sorted(lei.items(), key=lambda kv: -kv[1])

    maior = max(match_raw, key=lambda r: parse_brl(r[MATCH_VALOR_IDX])) if match_raw else None
    n_match, n_elab = len(match_raw), len(elab_raw)
    # "a apurar" = contratos no modelo de captacao (condicao contem "captado"); mais robusto
    # que contar Valor vazio (em julho os captados vem com Valor 0, nao vazio).
    elab_apurar = sum(1 for r in elab_raw
                      if "captado" in (r[ELAB_COND_IDX] if len(r) > ELAB_COND_IDX else "").lower())
    elab_fechado_n = n_elab - elab_apurar

    ext = ie.get("Externo", (0, 0.0))
    inte = ie.get("Interno", (0, 0.0))
    esf_clause = "; ".join(f"{k} com R$ {fmt_brl(v[1])} ({v[0]})" for k, v in esf_rows)

    if n_match == 0:
        match_sent = (
            f"No ciclo de {periodo} de 2026, não houve conversão de vendas de MATCH no período. "
            f"A Brada registrou {n_elab} contratos de Elaboração de Projetos. ")
    else:
        match_sent = (
            f"No ciclo de {periodo} de 2026, a Brada registrou {n_match} vendas de MATCH, "
            f"somando R$ {fmt_brl(match_total)}, e {n_elab} contratos de Elaboração de Projetos. "
            f"Em MATCH, {ext[0]} vendas foram externas (R$ {fmt_brl(ext[1])}) e {inte[0]} internas "
            f"(R$ {fmt_brl(inte[1])}). Por esfera de incentivo: {esf_clause}. "
            + (f"O maior aporte do período foi {maior[0]} (R$ {fmt_brl(parse_brl(maior[MATCH_VALOR_IDX]))}). "
               if maior else ""))
    fechado_txt = (f"{elab_fechado_n} contrato tem" if elab_fechado_n == 1
                   else f"{elab_fechado_n} contratos têm")
    apurar_txt = f"{elab_apurar} segue" if elab_apurar == 1 else f"{elab_apurar} seguem"
    exec_txt = (
        match_sent
        + f"Na Elaboração, {fechado_txt} valor fechado (R$ {fmt_brl(elab_total)}) "
        f"e {apurar_txt} no modelo de 10% sobre o valor captado, ainda a apurar. "
        "As comissões e os valores de pagamento são apurados pelo Financeiro."
    )

    cons_ie = [[k, str(v[0]), "R$ " + fmt_brl(v[1])] for k, v in ie_rows]
    cons_ie.append(["Total", str(n_match), "R$ " + fmt_brl(match_total)])
    cons_esf = [[k, str(v[0]), "R$ " + fmt_brl(v[1])] for k, v in esf_rows]
    cons_esf.append(["Total", str(n_match), "R$ " + fmt_brl(match_total)])
    cons_lei = [[k, str(v)] for k, v in lei_rows]
    cons_lei.append(["Total", str(n_elab)])

    return dict(exec_txt=exec_txt, cons_ie=cons_ie, cons_esf=cons_esf, cons_lei=cons_lei,
                elab_apurar=elab_apurar, elab_fechado_n=elab_fechado_n)


# ------------------------------------------------------------------ helpers docx
def _set_cell(cell, text, bold=False, color=DARK, size=8.0, align=None, shade=None):
    if shade:
        sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear")
        sh.set(qn("w:color"), "auto")
        sh.set(qn("w:fill"), shade)
        cell._element.get_or_add_tcPr().append(sh)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ORANGE if level <= 2 else DARK
        run.font.name = "Calibri"
    h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    h.paragraph_format.space_after = Pt(4)
    return h


def _label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = DARK
    r.bold = True
    return p


def _para(doc, text, size=10.5, color=DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def _note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.color.rgb = LIGHT
    r.italic = True


def _hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), HEADER_BG)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_table(doc, headers, rows, widths, aligns, bold_last=False):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    borders = OxmlElement("w:tblBorders")
    for bn in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{bn}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "D9D9D9")
        borders.append(b)
    tblPr.append(borders)
    for i, h in enumerate(headers):
        _set_cell(table.rows[0].cells[i], h, bold=True, color=WHITE, size=8.0,
                  align=aligns[i], shade=HEADER_BG)
    last = len(rows) - 1
    for r_idx, row in enumerate(rows):
        is_total = bold_last and r_idx == last
        shade = TOTAL_BG if is_total else (ROW_ALT_BG if r_idx % 2 == 1 else None)
        for c_idx, val in enumerate(row):
            _set_cell(table.rows[r_idx + 1].cells[c_idx], val, size=8.0, bold=is_total,
                      align=aligns[c_idx], shade=shade)
    for i, w in enumerate(widths):
        table.columns[i].width = Cm(w)
        for row in table.rows:
            row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


# ------------------------------------------------------------------ helpers html
def _css_align(a):
    if a == C:
        return "center"
    if a == R:
        return "right"
    return "left"


def _esc(s):
    return _htmllib.escape(s or "")


def _html_table(headers, rows, widths, aligns, cls="", bold_last=False):
    total = sum(widths)
    cols = "".join(f'<col style="width:{w / total * 100:.2f}%">' for w in widths)
    ths = "".join(f'<th style="text-align:{_css_align(aligns[i])}">{_esc(h)}</th>'
                  for i, h in enumerate(headers))
    last = len(rows) - 1
    body = []
    for r_idx, row in enumerate(rows):
        trcls = ' class="tot"' if (bold_last and r_idx == last) else ""
        tds = "".join(f'<td style="text-align:{_css_align(aligns[i])}">{_esc(v)}</td>'
                      for i, v in enumerate(row))
        body.append(f"<tr{trcls}>{tds}</tr>")
    c = f' class="{cls}"' if cls else ""
    return (f'<table{c}><colgroup>{cols}</colgroup><thead><tr>{ths}</tr></thead>'
            f'<tbody>{"".join(body)}</tbody></table>')


def render_html(path, emissao, summary, match_rows, elab_rows, ciclo=CICLO):
    has_match = bool(match_rows)
    css = """
@page { size: A4 landscape; margin: 1.2cm; }
* { box-sizing: border-box; }
body { font-family: Calibri, 'Segoe UI', Arial, sans-serif; color:#333; font-size:8pt;
       margin:0; -webkit-print-color-adjust:exact; print-color-adjust:exact; }
h1 { color:#C55A11; font-size:22pt; margin:0 0 2px 0; }
.sub { font-size:12pt; font-weight:bold; color:#333; margin:0 0 2px 0; }
.meta { font-size:10pt; color:#666; margin:0 0 8px 0; }
hr { border:none; border-top:2px solid #C55A11; margin:8px 0; }
h2 { color:#C55A11; font-size:13pt; margin:12px 0 4px 0; }
p.exec { font-size:10.5pt; color:#333; text-align:justify; margin:2px 0 8px 0; line-height:1.35; }
.cons-wrap { display:flex; flex-wrap:wrap; gap:22px; margin-bottom:6px; }
.cons-block h3 { color:#333; font-size:9.5pt; margin:2px 0 3px 0; }
table { border-collapse:collapse; width:100%; table-layout:fixed; margin-bottom:2px; }
thead { display:table-header-group; }
tr { break-inside:avoid; page-break-inside:avoid; }
h2, h3 { break-after:avoid; page-break-after:avoid; }
th { background:#C55A11; color:#fff; font-size:8pt; padding:3px 4px; border:1px solid #D9D9D9;
     font-weight:bold; word-wrap:break-word; overflow-wrap:break-word; vertical-align:top; }
td { font-size:8pt; padding:3px 4px; border:1px solid #D9D9D9; word-wrap:break-word;
     overflow-wrap:anywhere; vertical-align:top; }
tbody tr:nth-child(even) td { background:#FBE5D6; }
tbody tr.tot td { font-weight:bold; background:#F2E2D6; }
table.mini { width:auto; table-layout:auto; }
table.mini th, table.mini td { white-space:nowrap; padding:3px 12px 3px 6px; }
.note { font-size:9pt; color:#666; font-style:italic; margin:2px 0 10px 0; }
"""
    cons_blocks = []
    if has_match:
        cons_blocks += [
            "<div class='cons-block'><h3>MATCH por tipo</h3>",
            _html_table(["Tipo", "Negócios", "Valor"], summary["cons_ie"], [3, 1.4, 2.4],
                        [L, R, R], cls="mini", bold_last=True), "</div>",
            "<div class='cons-block'><h3>MATCH por esfera de incentivo</h3>",
            _html_table(["Esfera", "Negócios", "Valor"], summary["cons_esf"], [3.4, 1.4, 2.4],
                        [L, R, R], cls="mini", bold_last=True), "</div>",
        ]
    cons_blocks += [
        "<div class='cons-block'><h3>Elaboração por lei</h3>",
        _html_table(["Lei da submissão", "Negócios"], summary["cons_lei"], [3.2, 1.4],
                    [L, R], cls="mini", bold_last=True), "</div>",
    ]
    parts = [
        "<!doctype html><html><head><meta charset='utf-8'><style>", css, "</style></head><body>",
        "<h1>Relatório de Vendas e Comissões</h1>",
        f"<p class='sub'>{_esc(ciclo)}</p>",
        f"<p class='meta'>Brada &middot; Emissão {_esc(emissao)} &middot; "
        "Fonte: HubSpot (planilha Comissões 2026)</p>",
        "<hr>",
        "<h2>Sumário executivo</h2>",
        f"<p class='exec'>{_esc(summary['exec_txt'])}</p>",
        "<h2>Visão consolidada</h2>",
        "<div class='cons-wrap'>", *cons_blocks, "</div>",
        "<hr>",
        "<h2>Detalhamento das vendas MATCH</h2>",
    ]
    if has_match:
        parts += [
            _html_table(MATCH_HEADERS, match_rows, MATCH_WIDTHS, MATCH_ALIGN),
            "<p class='note'>As colunas de comissão são preenchidas pelo Financeiro.</p>",
        ]
        cobr = cobranca_rows(match_rows)
        if cobr:
            parts += [
                "<h2>MATCH externo (para cobrança)</h2>",
                _html_table(COBR_HEADERS, cobr, COBR_WIDTHS, COBR_ALIGN),
                "<p class='note'>Externos que o Financeiro precisa cobrar no período.</p>",
            ]
    else:
        parts += [f"<p class='note'>{_esc(NOTA_MATCH_VAZIO)}</p>"]
    parts += [
        "<h2>Detalhamento da Elaboração de Projetos</h2>",
        _html_table(ELAB_HEADERS, elab_rows, ELAB_WIDTHS, ELAB_ALIGN),
        "<p class='note'>As colunas de OBS, líquido pago e comissão são preenchidas pelo Financeiro.</p>",
        "</body></html>",
    ]
    Path(path).write_text("".join(parts), encoding="utf-8")


# Candidatos de navegador, em ordem de preferencia, para o PATH.
NAVEGADORES = ("google-chrome", "google-chrome-stable", "chromium", "chromium-browser", "chrome")


def achar_navegador():
    """Onde esta o navegador que imprime o PDF, ou None.

    A rota de PDF e "html -> navegador headless" justamente por ser portavel. Hardcodar o
    caminho do Windows anulava isso e era uma das tres amarras que impediam o script de sair
    da maquina do Bruno.

    Ordem: CHROME_BIN, depois o PATH, depois os lugares padrao do Windows.
    """
    do_env = os.environ.get("CHROME_BIN")
    if do_env and Path(do_env).exists():
        return do_env
    for nome in NAVEGADORES:
        achado = shutil.which(nome)
        if achado:
            return achado
    for caminho in (WIN_CHROME_1, WIN_CHROME_2):
        if Path(caminho).exists():
            return caminho
    return None


def html_to_pdf(html_path, pdf_path):
    navegador = achar_navegador()
    if not navegador:
        print("  [pdf] nenhum navegador encontrado. Instale o Chrome/Chromium ou aponte "
              "CHROME_BIN. O .md e o .docx foram gerados assim mesmo.")
        return False
    if Path(pdf_path).exists():
        Path(pdf_path).unlink()
    profile = Path(tempfile.gettempdir()) / "chrome_relatorio_profile"
    url = "file:///" + str(Path(html_path).resolve()).replace("\\", "/")
    cmd = [navegador, "--headless=new", "--disable-gpu", "--no-pdf-header-footer",
           f"--user-data-dir={profile}", f"--print-to-pdf={pdf_path}", url]
    subprocess.run(cmd, capture_output=True, timeout=120)
    return Path(pdf_path).exists()


# ------------------------------------------------------------------ render md
def render_md(path, emissao, summary, match_rows, elab_rows, ciclo=CICLO):
    has_match = bool(match_rows)

    def em(s):
        return (s or "").replace("|", "\\|")

    def md_table(headers, rows):
        out = ["| " + " | ".join(em(h) for h in headers) + " |",
               "|" + "|".join([" --- "] * len(headers)) + "|"]
        for r in rows:
            out.append("| " + " | ".join(em(c) for c in r) + " |")
        return "\n".join(out)

    lines = [
        "# Relatório de Vendas e Comissões", "",
        f"**{ciclo}**", "",
        f"Brada · Emissão {emissao} · Fonte: HubSpot (planilha Comissões 2026)", "",
        "## Sumário executivo", "", summary["exec_txt"], "",
        "## Visão consolidada", "",
    ]
    if has_match:
        lines += [
            "**MATCH por tipo**", "", md_table(["Tipo", "Negócios", "Valor"], summary["cons_ie"]), "",
            "**MATCH por esfera de incentivo**", "",
            md_table(["Esfera", "Negócios", "Valor"], summary["cons_esf"]), "",
        ]
    lines += [
        "**Elaboração por lei**", "",
        md_table(["Lei da submissão", "Negócios"], summary["cons_lei"]), "",
        "## Detalhamento das vendas MATCH", "",
    ]
    if has_match:
        lines += [
            md_table(MATCH_HEADERS, match_rows), "",
            "_As colunas de comissão são preenchidas pelo Financeiro._", "",
        ]
        cobr = cobranca_rows(match_rows)
        if cobr:
            lines += [
                "## MATCH externo (para cobrança)", "",
                md_table(COBR_HEADERS, cobr), "",
                "_Externos que o Financeiro precisa cobrar no período._", "",
            ]
    else:
        lines += [f"_{NOTA_MATCH_VAZIO}_", ""]
    lines += [
        "## Detalhamento da Elaboração de Projetos", "",
        md_table(ELAB_HEADERS, elab_rows), "",
        "_As colunas de OBS, líquido pago e comissão são preenchidas pelo Financeiro._", "",
    ]
    Path(path).write_text("\n".join(lines), encoding="utf-8")


# ------------------------------------------------------------------ render docx
def render_docx(path, emissao, summary, match_rows, elab_rows, ciclo=CICLO):
    has_match = bool(match_rows)
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)   # A4 paisagem (padrao BR), nao Letter (default do python-docx)
    sec.page_height = Cm(21.0)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)

    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.font.color.rgb = DARK
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.1

    title = doc.add_heading("Relatório de Vendas e Comissões", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = ORANGE
        run.font.name = "Calibri"
        run.font.size = Pt(22)

    sub = doc.add_paragraph()
    r = sub.add_run(ciclo)
    r.font.name = "Calibri"; r.font.size = Pt(12); r.font.color.rgb = DARK; r.bold = True
    sub.paragraph_format.space_after = Pt(2)

    meta = doc.add_paragraph()
    r = meta.add_run(f"Brada · Emissão {emissao} · Fonte: HubSpot (planilha Comissões 2026)")
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.color.rgb = LIGHT
    meta.paragraph_format.space_after = Pt(2)

    _hr(doc)

    _heading(doc, "Sumário executivo", 1)
    _para(doc, summary["exec_txt"])

    _heading(doc, "Visão consolidada", 1)
    if has_match:
        _label(doc, "MATCH por tipo")
        add_table(doc, ["Tipo", "Negócios", "Valor"], summary["cons_ie"],
                  [6.0, 3.0, 4.0], [L, R, R], bold_last=True)
        _label(doc, "MATCH por esfera de incentivo")
        add_table(doc, ["Esfera", "Negócios", "Valor"], summary["cons_esf"],
                  [7.0, 3.0, 4.0], [L, R, R], bold_last=True)
    _label(doc, "Elaboração por lei")
    add_table(doc, ["Lei da submissão", "Negócios"], summary["cons_lei"],
              [7.0, 3.0], [L, R], bold_last=True)

    _hr(doc)

    _heading(doc, "Detalhamento das vendas MATCH", 1)
    if has_match:
        add_table(doc, MATCH_HEADERS, match_rows, MATCH_WIDTHS, MATCH_ALIGN)
        _note(doc, "As colunas de comissão são preenchidas pelo Financeiro.")
        cobr = cobranca_rows(match_rows)
        if cobr:
            _heading(doc, "MATCH externo (para cobrança)", 1)
            add_table(doc, COBR_HEADERS, cobr, COBR_WIDTHS, COBR_ALIGN)
            _note(doc, "Externos que o Financeiro precisa cobrar no período.")
    else:
        _note(doc, NOTA_MATCH_VAZIO)

    _heading(doc, "Detalhamento da Elaboração de Projetos", 1)
    add_table(doc, ELAB_HEADERS, elab_rows, ELAB_WIDTHS, ELAB_ALIGN)
    _note(doc, "As colunas de OBS, líquido pago e comissão são preenchidas pelo Financeiro.")

    doc.save(path)


# ------------------------------------------------------------------ main
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--emissao", default=datetime.date.today().strftime("%d/%m/%Y"),
                    help="Data de emissao (dd/mm/aaaa ou aaaa-mm-dd). Default: hoje.")
    ap.add_argument("--cycle", default=None,
                    help="Ciclo YYYY-MM (ex.: 2026-07 = 20/06 a 20/07). Default: junho (tabs Junho_*).")
    ap.add_argument("--out", default=DEFAULT_OUT,
                    help="pasta de saida. Tambem sai de RELATORIOS_OUT.")
    ap.add_argument("--sheet-id", default=SID, help="planilha de origem (default: Comissoes 2026)")
    args = ap.parse_args()
    utf8_stdout()

    emissao = args.emissao
    if "-" in emissao and len(emissao) == 10:  # aaaa-mm-dd -> dd/mm/aaaa
        y, m, d = emissao.split("-")
        emissao = f"{d}/{m}/{y}"

    if args.cycle:
        if not re.fullmatch(r"\d{4}-(0[1-9]|1[0-2])", args.cycle):
            raise SystemExit(f"--cycle invalido: {args.cycle!r} (esperado YYYY-MM)")
        lab = cycle_labels(args.cycle)
        match_tab, elab_tab = lab["match_tab"], lab["elab_tab"]
        ciclo, periodo, basename = lab["ciclo"], lab["periodo"], lab["basename"]
    else:
        match_tab, elab_tab = MATCH_TAB, ELAB_TAB
        ciclo, periodo, basename = CICLO, PERIODO, BASENAME

    # Credencial pela cascata do repo (env JSON no CI, arquivo em prod, ~/.brada-secrets no dev),
    # e nao por caminho fixo. Era a segunda amarra de maquina.
    sh = get_sheets_client().open_by_key(args.sheet_id)

    match_rows, match_raw, match_total, _ = read_tab(
        sh, match_tab, MATCH_HEADERS, MATCH_BLANK, MATCH_MONEY, MATCH_TOTAL_IDX)
    elab_rows, elab_raw, elab_total, elab_apurar = read_tab(
        sh, elab_tab, ELAB_HEADERS, ELAB_BLANK, ELAB_MONEY, ELAB_TOTAL_IDX)

    summary = build_summary(match_raw, elab_raw, match_total, elab_total, elab_apurar, periodo)

    outdir = Path(args.out)
    outdir.mkdir(parents=True, exist_ok=True)
    md_path = outdir / f"{basename}.md"
    docx_path = outdir / f"{basename}.docx"
    pdf_path = outdir / f"{basename}.pdf"
    html_path = Path(tempfile.gettempdir()) / f"{basename}.html"

    render_md(md_path, emissao, summary, match_rows, elab_rows, ciclo)
    render_docx(docx_path, emissao, summary, match_rows, elab_rows, ciclo)
    render_html(html_path, emissao, summary, match_rows, elab_rows, ciclo)
    pdf_ok = html_to_pdf(html_path, pdf_path)

    print(f"Ciclo: {ciclo} | tabs: {match_tab} / {elab_tab}")
    print(f"MATCH: {len(match_rows)} linhas | total R$ {fmt_brl(match_total)}")
    print(f"Elaboração: {len(elab_rows)} linhas | fechado R$ {fmt_brl(elab_total)} | a apurar {summary['elab_apurar']}")
    print(f"MD:   {md_path}")
    print(f"DOCX: {docx_path}")
    print(f"PDF:  {pdf_path} ({'OK' if pdf_ok else 'FALHOU'})")


if __name__ == "__main__":
    main()
